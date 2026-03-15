"""
Word文档导出模块
将Markdown计算书转换为Word格式 - 支持LaTeX公式简化
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def create_word_report(md_text: str) -> bytes:
    """将 Markdown 计算书转换为 Word 二进制流"""
    doc = Document()

    # 设置全文字体样式
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    doc.add_heading('自承式管线桥结构设计计算书', 0)

    lines = md_text.split('\n')
    in_table = False
    table_data = []

    def flush_table():
        nonlocal in_table, table_data
        if table_data:
            valid_rows = [r for r in table_data if not re.match(r'^[\s\|-]*$', r)]
            if valid_rows:
                try:
                    col_count = len([c for c in valid_rows[0].split('|') if c.strip()])
                    table = doc.add_table(rows=len(valid_rows), cols=col_count)
                    table.style = 'Table Grid'
                    for i, row in enumerate(valid_rows):
                        cols = [c.strip() for c in row.split('|')[1:-1]]
                        for j, text in enumerate(cols):
                            if j < col_count:
                                # 清理公式
                                cell_text = clean_latex(text)
                                table.cell(i, j).text = cell_text
                except:
                    pass
            table_data = []
        in_table = False

    def clean_latex(text: str) -> str:
        """清理LaTeX公式，保留可读文本"""
        if not text:
            return text
        
        # 处理行内公式 $...$
        text = re.sub(r'\$([^$]+)\$', r'\1', text)
        
        # 处理行间公式 $$...$$
        text = re.sub(r'\$\$([^$]+)\$\$', r'\1', text)
        
        # 替换常见LaTeX命令为可读文本
        replacements = [
            (r'\\frac\{([^\}]+)\}\{([^\}]+)\}', r'(\1)/(\2)'),
            (r'\\sqrt\{([^\}]+)\}', r'sqrt(\1)'),
            (r'\\pi', 'π'),
            (r'\\alpha', 'α'),
            (r'\\beta', 'β'),
            (r'\\gamma', 'γ'),
            (r'\\delta', 'δ'),
            (r'\\theta', 'θ'),
            (r'\\rho', 'ρ'),
            (r'\\sigma', 'σ'),
            (r'\\phi', 'φ'),
            (r'\\psi', 'ψ'),
            (r'\\omega', 'ω'),
            (r'\\mu', 'μ'),
            (r'\\epsilon', 'ε'),
            (r'\\tau', 'τ'),
            (r'\\cdot', '·'),
            (r'\\times', '×'),
            (r'\\div', '÷'),
            (r'\\approx', '≈'),
            (r'\\leq', '≤'),
            (r'\\geq', '≥'),
            (r'\\neq', '≠'),
            (r'\\pm', '±'),
            (r'\\rightarrow', '→'),
            (r'\\leftarrow', '←'),
            (r'\\infty', '∞'),
            (r'\\Delta', 'Δ'),
            (r'\\Omega', 'Ω'),
            (r'\\lambda', 'λ'),
            (r'\\cdot', '·'),
            (r'\\left\(', '('),
            (r'\\right\)', ')'),
            (r'\\left\[', '['),
            (r'\\right\]', ']'),
            (r'\\left\{', '{'),
            (r'\\right\}', '}'),
            (r'\\/', ''),
            (r'\\', ''),
            (r'\{', ''),
            (r'\}', ''),
        ]
        
        for pattern, replacement in replacements:
            text = re.sub(pattern, replacement, text)
        
        return text

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 处理表格
        if line.startswith('|'):
            in_table = True
            table_data.append(line)
            continue
        else:
            if in_table:
                flush_table()

        # 处理标题
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('* **') or line.startswith('- **'):
            clean_text = clean_latex(line.replace('* **', '').replace('- **', '').replace('**', ''))
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(clean_text)
        elif line.startswith('---'):
            doc.add_paragraph()
        else:
            # 普通段落 - 清理公式但保留可读性
            clean_text = clean_latex(line)
            if clean_text:
                p = doc.add_paragraph()
                # 处理加粗 **text**
                parts = re.split(r'\*\*([^*]+)\*\*', clean_text)
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        # 加粗部分
                        run = p.add_run(part)
                        run.bold = True
                    elif part:
                        p.add_run(part)

    if in_table:
        flush_table()

    # 将文档保存到内存字节流
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
